#!/usr/bin/python
#coding:utf-8
#@ZHOU_YING
#2020-5-21
#提取报告中的信息

import os
import pandas as pd
from collections import defaultdict
import docx
from docx import Document

rootpath='/mnt/d/xinhua/2020.5.20_DATA/2000001B-3000000B/2000001B-2000500B/'
os.chdir(rootpath)
def dir_name(root):
#提取子文件夹名
    #os.chdir(root)
    for roots, dirs, files in os.walk(root):
        if roots == root:
            return dirs
            
def report_name(path):
#寻找到二代测序的word报告
    for roots,dirs,files in os.walk(path):
        if roots != path:
            break
        else:
            for file in files:
                if os.path.splitext(file)[1]=='.docx' and file.split('_')[0]=='NGS':
                    return file
            

dirslist=dir_name(rootpath)
DataInfo=defaultdict(list) #initial table

for m in range(55):
    #获取filename
    filename=dirslist[m]+'/'+report_name(dirslist[m])
    #print(filename)
    '''提取word报告中的信息'''
    doc=Document(filename)
    tables=doc.tables
    paragraph=doc.paragraphs

    '''样本信息提取'''
    table=tables[0]
    for i in range(0,len(table.rows)):
        table_row=table.rows[i]
        if i ==0:
            for cell in table_row.cells:
                if cell.text=='':
                    DataInfo[temp[0]].append('NA')
                else:
                    temp=cell.text.split('：')
                    if len(temp)==1:
                        temp=cell.text.split(':')
                    DataInfo[temp[0]].append(temp[1])
                    
        elif i==1 or i==2:
            for n in range(len(table_row.cells)-1):
                temp=table_row.cells[n].text.split('：')
                if len(temp)==1:
                    temp=table_row.cells[n].split(':')
                DataInfo[temp[0]].append(temp[1])
        else:
            DataInfo[table_row.cells[0].text].append(table_row.cells[1].text)
        
    '''检测项目信息'''
    table=tables[1]
    for i in range(0,len(table.rows)):
        table_row=table.rows[i]
        DataInfo[table_row.cells[0].text].append(table_row.cells[1].text)
    
    '''检测结果'''
    row_num=0
    for para in paragraph:
        if '3. 检测结果' in para.text:
            #print(paragraph[row_num+1].text)
            DataInfo['检测结果'].append(paragraph[row_num+1].text)
        row_num+=1
    '''有变异位点表格的'''
    table=tables[2]
    if '染色体' in table.rows[0].cells[0].text:
        row_num=len(table.rows)
        cell_num=len(table.rows[0].cells)
        for i in range(0,cell_num):
            info=''
            for j in range(1,row_num):
                info+=table.rows[j].cells[i].text+'; '
            code='DataInfo[table.rows[0].cells['+str(i)+'].text].append(info)'
            exec(code)
    else:
        m=['染色体','位置(hg19)','基因','变异','来源','疾病名称及遗传方式','致病性评级']
        for i in range(len(m)):
            code='DataInfo[m['+str(i)+']].append(\'NA\')'
            exec(code)

    '''结果说明及建议'''
    #前面有变异位点表格的
    if '染色体' in table.rows[0].cells[0].text:
        table=tables[3]
        #print(table.rows[0].cells[0].text)
        DataInfo['结果说明及建议'].append(table.rows[0].cells[0].text)
    #前面没有变异位点表格
    else:
        table=tables[2]
        DataInfo['结果说明及建议'].append(table.rows[0].cells[0].text)

#存储表格
os.chdir(rootpath)
DataInfo=pd.DataFrame(DataInfo)
DataInfo.to_csv('DataInfo.csv',index=False,encoding='utf_8_sig')