#!/usr/bin/python
#coding:utf-8
#@ZHOU_YING
#2020-5-21
#提取报告中的信息v1

import os
import pandas as pd
from collections import defaultdict
import docx
from docx import Document

rootpath='/data/zhouying/tmp_doc/'
os.chdir(rootpath)

dirslist=os.listdir()
DataInfo=pd.DataFrame() #initial table

for m in range(len(dirslist)):
    #获取filename
    filename=dirslist[m]
    print(str(m)+'. '+filename+'  working...')
    '''提取word报告中的信息'''
    inform=defaultdict(list)
    doc=Document(filename)
    tables=doc.tables
    paragraph=doc.paragraphs
    txt=''
    for para in paragraph:
        txt+=para.text
    if '高通量测序检测报告' not in txt:
        continue
    '''样本信息提取'''
    table=tables[0]
    for i in range(0,len(table.rows)):
        back=''
        table_row=table.rows[i]
        if i ==0:
            for cell in table_row.cells:
                if cell.text=='':
                    inform[temp[0]].append('NA')
                elif cell.text!=back or back=='':
                    temp=cell.text.split('：')
                    if len(temp)==1:
                        temp=cell.text.split(':')
                    inform[temp[0]].append(temp[1])
                else:
                    continue
                back=cell.text
                    
        elif i==1 or i==2:
            for n in range(len(table_row.cells)-1):
                temp=table_row.cells[n].text.split('：')
                if len(temp)==1:
                    temp=table_row.cells[n].text.split(':')
                inform[temp[0]].append(temp[1])
        else:
            inform[table_row.cells[0].text].append(table_row.cells[1].text)
        
    '''检测项目信息'''
    table=tables[1]
    for i in range(0,len(table.rows)):
        table_row=table.rows[i]
        inform[table_row.cells[0].text].append(table_row.cells[1].text)
    
    '''检测结果'''
    row_num=0
    temp1=''
    for para in paragraph:
        if '3. 检测结果' in para.text:
            n=0
            while '4.' not in paragraph[row_num+n].text:
                n+=1
                temp1+=paragraph[row_num+n].text+'\n'
                #print(temp1)
            
        row_num+=1
        temp1=temp1.replace('\n4. 结果说明及建议\n','')
    inform['检测结果'].append(temp1)
    '''有变异位点表格的'''
    table=tables[2]
    if '染色体' in table.rows[0].cells[0].text and '2' not in table.rows[0].cells[0].text:
        row_num=len(table.rows)
        cell_num=len(table.rows[0].cells)
        for i in range(0,cell_num):
            info=''
            for j in range(1,row_num):
                info+=table.rows[j].cells[i].text+'|'
            code='inform[table.rows[0].cells['+str(i)+'].text].append(info)'
            exec(code)
    else:
        m=['染色体','位置(hg19)','基因','变异','来源','疾病名称及遗传方式','致病性评级']
        for i in range(len(m)):
            code='inform[m['+str(i)+']].append(\'NA\')'
            exec(code)

        '''结果说明及建议'''
        #前面有变异位点表格的
    if '染色体' in table.rows[0].cells[0].text and '2' not in table.rows[0].cells[0].text:
        table=tables[3]
        #print(table.rows[0].cells[0].text)
        inform['结果说明及建议'].append(table.rows[0].cells[0].text)
        #前面没有变异位点表格
    else:
        table=tables[2]
        inform['结果说明及建议'].append(table.rows[0].cells[0].text)
    inform=pd.DataFrame(inform)
    #inform.to_csv(filename+'.csv',index=False,encoding='utf_8_sig')
    DataInfo=pd.concat([DataInfo,inform],axis=0)

#存储表格
os.chdir(rootpath)
DataInfo=pd.DataFrame(DataInfo)
DataInfo.to_csv('DataInfo.csv',index=False,encoding='utf_8_sig')
